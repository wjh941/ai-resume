from __future__ import annotations

from pathlib import Path

from docx import Document

from app.schemas.resume import ResumePayload


def render_word_resume(
    resume: ResumePayload, output_path: Path, watermark_text: str | None = None
) -> None:
    document = Document()
    visibility = resume.section_visibility
    if visibility.basic:
        document.add_heading(resume.basic.name, level=0)
        document.add_paragraph(
            " | ".join(part for part in [resume.basic.phone, resume.basic.email, resume.basic.city] if part)
        )
    if visibility.job:
        document.add_heading("求职意向", level=1)
        document.add_paragraph(resume.job.target_role)
    if visibility.education:
        _add_entries(
            document,
            "教育经历",
            [
                f"{item.school} | {item.major} | {item.degree} | {item.start_date} - {item.end_date}"
                for item in resume.education
            ],
        )
    if visibility.employment:
        _add_entries(
            document,
            "实习/工作经历",
            [
                "\n".join(
                    part
                    for part in [
                        f"{item.company} | {item.position} | {item.start_date} - {item.end_date}",
                        item.description,
                    ]
                    if part
                )
                for item in resume.employment
            ],
        )
    if visibility.projects:
        _add_entries(
            document,
            "项目经历",
            [
                "\n".join(
                    part
                    for part in [
                        f"{item.name} | {item.role} | {item.start_date} - {item.end_date}",
                        item.description,
                    ]
                    if part
                )
                for item in resume.projects
            ],
        )
    if visibility.skills:
        _add_entries(document, "技能证书", [", ".join(resume.skills.skills + resume.skills.certificates)])
    if visibility.self_evaluation:
        _add_entries(document, "自我评价", [resume.self_evaluation])
    if watermark_text:
        watermark = document.add_paragraph(watermark_text)
        watermark.style = document.styles["Caption"]
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _add_entries(document: Document, heading: str, entries: list[str]) -> None:
    meaningful_entries = [entry for entry in entries if entry.strip()]
    if not meaningful_entries:
        return
    document.add_heading(heading, level=1)
    for entry in meaningful_entries:
        document.add_paragraph(entry, style="List Bullet")
