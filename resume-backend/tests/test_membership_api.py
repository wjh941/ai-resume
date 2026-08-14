from __future__ import annotations

from datetime import datetime, timedelta, timezone
from io import BytesIO
import sqlite3
import threading
from concurrent.futures import ThreadPoolExecutor

from docx import Document

from conftest import grant_vip, make_draft_payload, make_resume_payload

def _login_headers(api_client, phone: str) -> dict[str, str]:
    response = api_client.post(
        "/api/auth/login-phone", json={"phone": phone, "code": "123456"}
    )
    assert response.status_code == 200
    return {"Authorization": f"Bearer {response.json()['data']['token']}"}


def test_vip_info_rejects_missing_bearer_token(api_client):
    api_client.headers.pop("Authorization")

    response = api_client.get("/api/user/vip-info")

    assert response.status_code == 401


def test_new_user_receives_free_vip_status(api_client):
    response = api_client.get("/api/user/vip-info")

    assert response.status_code == 200
    assert response.json()["data"] == {
        "vip_level": "free",
        "expire_time": None,
        "auto_renew": False,
        "max_drafts": 3,
        "max_compare_jobs": 2,
    }


def test_demo_payment_fulfills_only_the_order_owner(api_client):
    packages = api_client.get("/api/pay/package-list")

    assert packages.status_code == 200
    assert [item["package_type"] for item in packages.json()["data"]["items"]] == [
        "monthly", "quarterly", "annual"
    ]

    created = api_client.post(
        "/api/pay/create-order", json={"package_type": "monthly", "auto_renew": False}
    )
    assert created.status_code == 200
    order_id = created.json()["data"]["order_id"]
    assert created.json()["data"]["payment_status"] == "pending"

    completed = api_client.post(
        "/api/pay/callback",
        json={"order_id": order_id, "payment_channel": "demo", "payment_status": "paid"},
    )
    assert completed.status_code == 200
    assert completed.json()["data"]["order"]["payment_status"] == "paid"
    assert completed.json()["data"]["vip"]["vip_level"] == "basic"
    assert completed.json()["data"]["vip"]["expire_time"] is not None

    other_orders = api_client.get(
        "/api/user/order-list", headers=_login_headers(api_client, "13900139000")
    )
    assert other_orders.status_code == 200
    assert other_orders.json()["data"]["items"] == []

    stolen_callback = api_client.post(
        "/api/pay/callback",
        headers=_login_headers(api_client, "13900139000"),
        json={"order_id": order_id, "payment_channel": "demo", "payment_status": "paid"},
    )
    assert stolen_callback.status_code == 404


def test_demo_payment_callback_is_idempotent(api_client):
    order = api_client.post("/api/pay/create-order", json={"package_type": "monthly"}).json()["data"]
    payload = {"order_id": order["order_id"], "payment_channel": "demo", "payment_status": "paid"}

    first = api_client.post("/api/pay/callback", json=payload)
    second = api_client.post("/api/pay/callback", json=payload)

    assert first.status_code == 200
    assert second.status_code == 200
    assert first.json()["data"]["order"]["entitlement_expire_time"] == second.json()["data"]["order"]["entitlement_expire_time"]


def test_unconfigured_real_payment_channel_leaves_order_pending(api_client):
    order = api_client.post("/api/pay/create-order", json={"package_type": "monthly"}).json()["data"]

    callback = api_client.post(
        "/api/pay/callback",
        json={"order_id": order["order_id"], "payment_channel": "wechat_pay", "payment_status": "paid"},
    )

    assert callback.status_code == 503
    assert callback.json()["code"] == "payment_channel_unavailable"
    assert api_client.get("/api/user/order-list").json()["data"]["items"][0]["payment_status"] == "pending"


def test_active_premium_user_cannot_purchase_a_lower_tier_package(api_client):
    grant_vip(api_client, "premium")

    response = api_client.post("/api/pay/create-order", json={"package_type": "monthly"})

    assert response.status_code == 409
    assert response.json()["code"] == "membership_package_conflict"
    assert api_client.get("/api/user/order-list").json()["data"]["items"] == []


def test_expired_vip_is_downgraded_when_its_status_is_read(api_client):
    with sqlite3.connect(api_client.app.state.settings.database_path) as connection:
        user_id = connection.execute(
            "SELECT user_id FROM users WHERE phone = ?", ("13800138000",)
        ).fetchone()[0]
        connection.execute(
            """
            INSERT INTO user_vip (user_id, vip_level, expire_time, auto_renew, create_time)
            VALUES (?, 'premium', ?, 1, ?)
            ON CONFLICT(user_id) DO UPDATE SET
                vip_level = excluded.vip_level,
                expire_time = excluded.expire_time,
                auto_renew = excluded.auto_renew
            """,
            (
                user_id,
                (datetime.now(timezone.utc) - timedelta(seconds=1)).isoformat(),
                datetime.now(timezone.utc).isoformat(),
            ),
        )

    response = api_client.get("/api/user/vip-info")

    assert response.status_code == 200
    assert response.json()["data"]["vip_level"] == "free"
    assert response.json()["data"]["expire_time"] is None


def test_free_user_cannot_create_a_fourth_resume_draft(api_client):
    for index in range(3):
        response = api_client.post(
            "/api/draft/save", json=make_draft_payload(job_title=f"Role {index}")
        )
        assert response.status_code == 200

    rejected = api_client.post(
        "/api/draft/save", json=make_draft_payload(job_title="Role 4")
    )

    assert rejected.status_code == 403
    assert rejected.json()["code"] == "vip_required"


def test_free_user_concurrent_draft_saves_cannot_exceed_the_limit(api_client, monkeypatch):
    for index in range(2):
        assert api_client.post(
            "/api/draft/save", json=make_draft_payload(job_title=f"Existing {index}")
        ).status_code == 200

    repository = api_client.app.state.draft_repository
    original_list = repository.list
    read_barrier = threading.Barrier(2)

    def synchronized_list(user_id: str):
        items = original_list(user_id)
        read_barrier.wait(timeout=5)
        return items

    monkeypatch.setattr(repository, "list", synchronized_list)
    with ThreadPoolExecutor(max_workers=2) as executor:
        responses = list(
            executor.map(
                lambda index: api_client.post(
                    "/api/draft/save", json=make_draft_payload(job_title=f"Concurrent {index}")
                ),
                range(2),
            )
        )

    assert sorted(response.status_code for response in responses) == [200, 403]
    assert len(original_list(api_client.app.state.auth_service.verify(api_client.headers["Authorization"].split(" ", 1)[1]))) == 3


def test_free_user_cannot_compare_more_than_two_jobs(api_client):
    profile = {
        "identity_code": "2",
        "major": "计算机科学与技术",
        "education_level": "本科",
        "graduation_year": 2027,
        "city_preferences": ["上海"],
        "minimum_salary": "10k",
        "industry_preferences": ["互联网"],
        "work_types": ["全职"],
        "skills": ["Python", "SQL"],
    }
    assert api_client.post("/api/career/profile/save", json=profile).status_code == 200

    rejected = api_client.post(
        "/api/career/compare",
        json={"role_names": ["数据工程师", "后端开发工程师", "前端开发工程师"]},
    )

    assert rejected.status_code == 403
    assert rejected.json()["code"] == "vip_required"


def test_free_word_export_contains_platform_watermark(api_client):
    draft = api_client.post("/api/draft/save", json=make_draft_payload()).json()["data"]

    exported = api_client.post(
        "/api/export/word", json={"draft_id": draft["id"]}
    )
    assert exported.status_code == 200
    downloaded = api_client.get(exported.json()["data"]["download_url"])

    document = Document(BytesIO(downloaded.content))
    content = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Resume Dashboard Free" in content


def test_basic_and_premium_word_exports_apply_the_correct_watermark(api_client):
    grant_vip(api_client, "basic")
    basic_draft = api_client.post("/api/draft/save", json=make_draft_payload()).json()["data"]
    basic_export = api_client.post("/api/export/word", json={"draft_id": basic_draft["id"]})
    basic_document = Document(BytesIO(api_client.get(basic_export.json()["data"]["download_url"]).content))
    assert "Resume Dashboard Basic" in "\n".join(paragraph.text for paragraph in basic_document.paragraphs)

    grant_vip(api_client, "premium")
    premium_draft = api_client.post("/api/draft/save", json=make_draft_payload()).json()["data"]
    premium_export = api_client.post("/api/export/word", json={"draft_id": premium_draft["id"]})
    premium_document = Document(BytesIO(api_client.get(premium_export.json()["data"]["download_url"]).content))
    premium_content = "\n".join(paragraph.text for paragraph in premium_document.paragraphs)
    assert "Resume Dashboard" not in premium_content


def test_production_configuration_rejects_demo_payment(tmp_path, monkeypatch):
    from fastapi.testclient import TestClient

    monkeypatch.setenv("DATABASE_PATH", str(tmp_path / "production.db"))
    monkeypatch.setenv("TEMP_FILE_PATH", str(tmp_path / "temp"))
    monkeypatch.setenv("APP_ENV", "production")
    monkeypatch.setenv("AUTH_DEMO_MODE", "true")
    monkeypatch.setenv("PAYMENT_DEMO_MODE", "true")
    monkeypatch.setenv("JWT_SECRET", "test-jwt-secret-for-production-payment")

    from main import create_app

    with TestClient(create_app()) as client:
        login = client.post("/api/auth/login-phone", json={"phone": "13800138000", "code": "123456"})
        client.headers.update({"Authorization": f"Bearer {login.json()['data']['token']}"})
        order = client.post("/api/pay/create-order", json={"package_type": "monthly"}).json()["data"]
        callback = client.post(
            "/api/pay/callback",
            json={"order_id": order["order_id"], "payment_channel": "demo", "payment_status": "paid"},
        )

    assert callback.status_code == 503
    assert callback.json()["code"] == "payment_demo_disabled"


def test_free_user_cannot_request_deep_ai_resume_polish(api_client):
    response = api_client.post(
        "/api/resume/ai-rewrite",
        json={
            "resume": make_resume_payload(),
            "job": {"version": 1, "role_name": "Data Engineer", "required_skills": ["Python"]},
            "mode": "deep",
        },
    )

    assert response.status_code == 403
    assert response.json()["code"] == "vip_required"


def test_free_assessment_response_excludes_full_career_roadmap(api_client):
    response = api_client.post(
        "/api/career/assessment/submit",
        json={"answers": {"interest_investigative_1": 5, "style_structure_1": 5}},
    )

    assert response.status_code == 200
    result = response.json()["data"]["result"]
    assert result["report_scope"] == "simplified"
    assert "action_plan" not in result


def test_free_career_recommendation_excludes_saved_assessment_roadmap(api_client):
    profile = {
        "identity_code": "2",
        "major": "Computer Science",
        "education_level": "Bachelor",
        "graduation_year": 2027,
        "city_preferences": ["Shanghai"],
        "minimum_salary": "10k",
        "industry_preferences": ["Internet"],
        "work_types": ["full-time"],
        "skills": ["Python", "SQL"],
    }
    assert api_client.post("/api/career/profile/save", json=profile).status_code == 200
    assert api_client.post(
        "/api/career/assessment/submit",
        json={"answers": {"interest_investigative_1": 5, "style_structure_1": 5}},
    ).status_code == 200

    recommendation = api_client.post("/api/career/recommend")

    assert recommendation.status_code == 200
    action_plan = recommendation.json()["data"]["assessment_guidance"]["action_plan"]
    assert all(not items for items in action_plan.values())


def test_free_user_cannot_open_full_job_report_or_industry_insights(api_client):
    job_report = api_client.post(
        "/api/consultation/job-analysis",
        json={"role_name": "Data Engineer", "identity_code": "2"},
    )
    insights = api_client.get("/api/career/annual-insights")

    assert job_report.status_code == 403
    assert job_report.json()["code"] == "vip_required"
    assert insights.status_code == 403
    assert insights.json()["code"] == "vip_required"
